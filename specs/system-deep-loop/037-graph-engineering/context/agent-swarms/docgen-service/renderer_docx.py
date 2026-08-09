"""
Server-side Word renderer (python-docx).

Consumes the same DocxPlan JSON the browser builder does:
  { "title": str,
    "blocks": [ {"type":"heading","level":1|2|3,"text":str}
              | {"type":"paragraph","text":str}
              | {"type":"bullets","items":[str]}
              | {"type":"table","table":{"columns":[str],"rows":[[cell]]}} ] }

Applies the Anthropic docx-skill guidance that the browser lib can't do as
cleanly natively:
  • a real cover page + an updatable Table of Contents field,
  • level-1 sections each starting on a fresh page (multi-page output),
  • fixed-layout tables with explicit column widths, a shaded header row
    (w:shd, not a black SOLID fill) and full gridlines,
  • proper bullet lists and generous paragraph spacing.

Everything degrades gracefully: a malformed block is skipped, never fatal.
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# Letter, 1" margins → 6.5" of usable width for full-bleed tables.
USABLE_WIDTH = Inches(6.5)
HEADER_FILL = "EEF2F7"       # light blue-grey, matches the browser builder
BORDER_COLOR = "D9DEE5"
ACCENT = RGBColor(0x37, 0x41, 0x51)


def _set_cell_shading(cell, fill: str) -> None:
    """Header cell background via w:shd (CLEAR pattern, NOT a black SOLID fill)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_width(cell, width) -> None:
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width.twips)))
    tcW.set(qn("w:type"), "dxa")


def _fixed_table_layout(table) -> None:
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _add_toc(doc) -> None:
    """A real TOC field. Word offers to update it (with page numbers) on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(placeholder)
    run._r.append(fldEnd)


def _cover(doc, title: str) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title or "Untitled")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d = sub.add_run(date.today().strftime("%B %d, %Y"))
    d.font.size = Pt(12)
    d.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_table(doc, table_spec: dict[str, Any]) -> None:
    columns = [str(c) for c in (table_spec.get("columns") or [])]
    rows = table_spec.get("rows") or []
    ncols = max(1, len(columns))
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"           # full gridlines
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    _fixed_table_layout(table)
    col_w = Inches(6.5 / ncols)

    hdr = table.rows[0].cells
    for i in range(ncols):
        _set_cell_width(hdr[i], col_w)
        _set_cell_shading(hdr[i], HEADER_FILL)
        cell_p = hdr[i].paragraphs[0]
        run = cell_p.add_run(columns[i] if i < len(columns) else "")
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            _set_cell_width(cells[i], col_w)
            val = row[i] if isinstance(row, list) and i < len(row) else None
            cells[i].paragraphs[0].add_run("" if val is None else str(val))
    doc.add_paragraph()


def render_docx(plan: dict[str, Any]) -> bytes:
    doc = Document()

    # Normal-style defaults: readable body font + line spacing.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    _cover(doc, str(plan.get("title") or "Untitled"))

    # TOC on its own page, then start the body on a fresh page.
    doc.add_page_break()
    heading = doc.add_paragraph()
    hr = heading.add_run("Contents")
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = ACCENT
    _add_toc(doc)
    doc.add_page_break()

    seen_h1 = False
    for block in plan.get("blocks") or []:
        if not isinstance(block, dict):
            continue
        btype = block.get("type")
        try:
            if btype == "heading":
                level = block.get("level", 2)
                level = level if level in (1, 2, 3) else 2
                if level == 1 and seen_h1:
                    doc.add_page_break()
                h = doc.add_heading(str(block.get("text") or ""), level=level)
                for run in h.runs:
                    run.font.color.rgb = ACCENT
                if level == 1:
                    seen_h1 = True
            elif btype == "paragraph":
                p = doc.add_paragraph(str(block.get("text") or ""))
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
            elif btype == "bullets":
                for item in block.get("items") or []:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif btype == "table":
                spec = block.get("table")
                if isinstance(spec, dict):
                    _add_table(doc, spec)
        except Exception:
            # One bad block never sinks the document.
            continue

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
